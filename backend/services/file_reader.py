import os
from typing import Optional

def read_text_from_file(file_path: str) -> Optional[str]:
    """Extract text content from a supported document file.

    Supports PDF, DOCX, and TXT. Returns None if unsupported or on failure.
    """
    if not file_path or not os.path.exists(file_path):
        return None

    _, ext = os.path.splitext(file_path.lower())

    try:
        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

        if ext == '.pdf':
            # Prefer lightweight extraction to avoid heavy dependencies
            try:
                import pdfplumber  # type: ignore
                text_parts = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text_parts.append(page.extract_text() or "")
                return "\n".join(text_parts).strip() or None
            except Exception:
                # Fallback to PyPDF2
                try:
                    import PyPDF2  # type: ignore
                    text_parts = []
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        for page in reader.pages:
                            text_parts.append(page.extract_text() or "")
                    return "\n".join(text_parts).strip() or None
                except Exception:
                    return None

        if ext == '.docx':
            try:
                import docx  # type: ignore
            except Exception:
                from docx import Document as _Doc  # type: ignore
                doc = _Doc(file_path)
                return "\n".join(p.text for p in doc.paragraphs).strip() or None

            from docx import Document  # type: ignore
            doc = Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs).strip() or None

        return None
    except Exception:
        return None


